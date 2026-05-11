"""
CV oluşturucu — Muhammed Sina Gün
Çalıştır: python3 generate_cv.py
Çıktı: muhammedsinagun_cv_tr.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Sayfa kenar boşlukları ──────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# ── Renk paleti ─────────────────────────────────────────────────────────────
DARK_NAVY  = RGBColor(0x1A, 0x1A, 0x2E)   # başlık koyu lacivert
ACCENT     = RGBColor(0x16, 0x21, 0x3E)   # bölüm çizgisi
MID_GRAY   = RGBColor(0x44, 0x44, 0x55)   # alt başlık
LIGHT_GRAY = RGBColor(0x66, 0x66, 0x77)   # yardımcı metin
BLACK      = RGBColor(0x1C, 0x1C, 0x1C)   # gövde metni

# ── Yardımcı fonksiyonlar ───────────────────────────────────────────────────

def set_font(run, size=10.5, bold=False, italic=False, color=BLACK,
             font_name="Calibri"):
    run.bold        = bold
    run.italic      = italic
    run.font.size   = Pt(size)
    run.font.color.rgb = color
    run.font.name   = font_name
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),    font_name)
    rFonts.set(qn("w:hAnsi"),   font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rPr.insert(0, rFonts)

def para_spacing(p, before=0, after=0, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def add_hr(doc, color_hex="1A1A2E", thickness=8):
    """Bölüm ayracı — tam genişlikte ince çizgi."""
    p  = doc.add_paragraph()
    para_spacing(p, before=2, after=4)
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(thickness))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pb.append(bot)
    pPr.append(pb)

def heading1(doc, text):
    """Ana bölüm başlığı."""
    p   = doc.add_paragraph()
    run = p.add_run(text.upper())
    set_font(run, size=11, bold=True, color=DARK_NAVY)
    para_spacing(p, before=10, after=2)
    add_hr(doc, color_hex="162136", thickness=6)

def heading2(doc, title, subtitle="", date_str=""):
    """İş / proje / okul başlığı."""
    p = doc.add_paragraph()
    para_spacing(p, before=6, after=1)
    r1 = p.add_run(title)
    set_font(r1, size=10.5, bold=True, color=DARK_NAVY)
    if subtitle:
        r2 = p.add_run(f"  |  {subtitle}")
        set_font(r2, size=10, italic=True, color=MID_GRAY)
    if date_str:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r3 = p.add_run(f"\n{date_str}")
        set_font(r3, size=9, italic=True, color=LIGHT_GRAY)

def bullet(doc, text, indent=0.4):
    """Madde işaretli satır."""
    p   = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run, size=10, color=BLACK)
    para_spacing(p, before=0, after=1)
    p.paragraph_format.left_indent  = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.18)

def body(doc, text, italic=False):
    """Normal gövde paragrafı."""
    p   = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=10, italic=italic, color=BLACK)
    para_spacing(p, before=1, after=2)

def tag_line(doc, text):
    """Teknoloji / etiket satırı."""
    p   = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=9.5, italic=True, color=LIGHT_GRAY)
    para_spacing(p, before=0, after=3)

# ═══════════════════════════════════════════════════════════════════════════
#  HEADER — Ad, unvan, iletişim
# ═══════════════════════════════════════════════════════════════════════════

p_name = doc.add_paragraph()
p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_name.add_run("MUHAMMEd SİNA GÜN")
set_font(r, size=22, bold=True, color=DARK_NAVY)
para_spacing(p_name, before=0, after=2)

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run(
    "Bilgisayar Mühendisliği Öğrencisi  |  "
    "Yapay Zekâ & Makine Öğrenmesi  |  "
    "Prompt Engineering  |  Yazılım Geliştirme"
)
set_font(r, size=10.5, italic=True, color=MID_GRAY)
para_spacing(p_title, before=0, after=4)

add_hr(doc, color_hex="1A1A2E", thickness=12)

p_contact = doc.add_paragraph()
p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spacing(p_contact, before=3, after=6)
contact_items = [
    ("Konum:", "İstanbul, Türkiye"),
    ("GitHub:", "github.com/msgxr"),
    ("LinkedIn:", "linkedin.com/in/muhammedsina"),
]
for i, (label, val) in enumerate(contact_items):
    rl = p_contact.add_run(label + " ")
    set_font(rl, size=9.5, bold=True, color=MID_GRAY)
    rv = p_contact.add_run(val)
    set_font(rv, size=9.5, color=BLACK)
    if i < len(contact_items) - 1:
        rs = p_contact.add_run("   ·   ")
        set_font(rs, size=9.5, color=LIGHT_GRAY)

# ═══════════════════════════════════════════════════════════════════════════
#  1. PROFESYONEL ÖZET
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "Profesyonel Özet")

ozet = (
    "İstanbul Arel Üniversitesi Bilgisayar Mühendisliği bölümünde lisans eğitimimi "
    "sürdürmekteyim. Akademik sürecimin yanı sıra yapay zekâ, makine öğrenmesi, "
    "Graph Neural Network mimarileri ve prompt engineering alanlarında uygulamalı "
    "proje çalışmaları yürütmekteyim. Özellikle sağlık alanında yapay zekâ "
    "uygulamaları üzerine geliştirmekte olduğum VARIANT-GNN projesi kapsamında; "
    "teknik dokümantasyon, hata analizi, model geliştirme süreçleri ve GitHub repo "
    "yönetimi konularında deneyim kazandım. Claude, ChatGPT ve benzeri büyük dil "
    "modellerini yazılım geliştirme, proje planlama, hata kontrol listesi hazırlama "
    "ve teknik çıktı üretimi amacıyla sistematik biçimde kullanmaktayım. Yazılım "
    "mühendisliği bakış açısıyla sistematik ve belgelenmiş çalışma anlayışını "
    "benimsemekteyim."
)
body(doc, ozet)

# ═══════════════════════════════════════════════════════════════════════════
#  2. TEKNİK YETKİNLİKLER
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "Teknik Yetkinlikler")

kategoriler = [
    (
        "Yapay Zekâ / Makine Öğrenmesi",
        [
            "Graph Neural Networks (GNN)",
            "Makine öğrenmesi model geliştirme ve değerlendirme süreçleri",
            "Açıklanabilir Yapay Zekâ (Explainable AI — XAI)",
            "Genetik varyant sınıflandırması üzerine uygulamalı çalışma",
            "Model çıktılarının yorumlanması ve analizi",
            "Yapay zekâ destekli proje geliştirme süreçleri",
        ],
    ),
    (
        "Prompt Engineering & Büyük Dil Modelleri (LLM)",
        [
            "Claude, ChatGPT ve benzeri LLM araçlarıyla uygulamalı çalışma",
            "Uzun ve yapılandırılmış master prompt hazırlama",
            "Teknik şartnameye uygun çıktı üretme promptları",
            "Hata kontrol listesi (checklist) promptları geliştirme",
            "Proje analiz, kod inceleme ve dokümantasyon iyileştirme promptları",
            "AI workflow entegrasyonu ve üretkenlik uygulamaları",
        ],
    ),
    (
        "Programlama Dilleri",
        [
            "Python — makine öğrenmesi, veri işleme, otomasyon (birincil dil)",
            "Java — nesne yönelimli yazılım geliştirme",
            "SQL — veritabanı sorgulama ve yönetimi",
            "C++ — algoritmik programlama",
            "JavaScript — temel düzey",
        ],
    ),
    (
        "Araçlar & Platformlar",
        [
            "Git & GitHub — sürüm kontrolü, repo yönetimi, branch stratejileri",
            "Streamlit — web tabanlı demo ve dashboard geliştirme",
            "Jupyter Notebook — veri bilimi ve model çalışmaları",
            "Docker — konteyner tabanlı geliştirme ortamı",
            "n8n — otomasyon iş akışı tasarımı",
            "Huawei Cloud — bulut servisleri ve proje altyapısı",
            "VS Code, PowerShell, Linux",
        ],
    ),
    (
        "Yazılım Mühendisliği Becerileri",
        [
            "Teknik dokümantasyon hazırlama ve düzenleme",
            "Hata analizi ve sistematik hata ayıklama süreci",
            "Gereksinim analizi ve yazılım modelleme",
            "GitHub üzerinden proje ve dosya yapısı yönetimi",
            "Proje planlama, süreç takibi ve otomasyon sistemi tasarımı",
        ],
    ),
    (
        "Kişisel ve Profesyonel Beceriler",
        [
            "Teknik detaylara odaklanma ve belgelenmiş çalışma alışkanlığı",
            "Karmaşık teknik konuları yazılı olarak ifade edebilme",
            "Araştırma odaklı ve sorgulayıcı öğrenme yaklaşımı",
            "Bağımsız proje yönetimi ve inisiyatif alma",
        ],
    ),
]

for kat_baslik, maddeler in kategoriler:
    p = doc.add_paragraph()
    r = p.add_run(kat_baslik)
    set_font(r, size=10, bold=True, color=ACCENT)
    para_spacing(p, before=5, after=1)
    for m in maddeler:
        bullet(doc, m)

# ═══════════════════════════════════════════════════════════════════════════
#  3. TEKNİK ÇALIŞMALAR & DENEYİM
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "Teknik Çalışmalar & Deneyim")

# --- Çalışma 1 ---
heading2(doc, "Yapay Zekâ Destekli Proje Geliştirme",
         "Bağımsız Çalışma", "İstanbul  |  Devam ediyor")
bullet(doc, "VARIANT-GNN projesi kapsamında sağlık alanında yapay zekâ uygulamaları üzerine çalışma yürütmekteyim.")
bullet(doc, "GNN mimarileri, genetik varyant sınıflandırması ve açıklanabilir yapay zekâ yöntemleri üzerine araştırma ve uygulama çalışmaları gerçekleştirdim.")
bullet(doc, "Proje sürecinde GitHub repo yönetimi, dosya ve kod yapısı düzenleme, teknik dokümantasyon hazırlama ve hata analizi süreçlerinde aktif olarak yer aldım.")
bullet(doc, "Streamlit ile demo ve dashboard yapısı üzerine çalışmalar yürüttüm.")

# --- Çalışma 2 ---
heading2(doc, "Prompt Engineering & AI Workflow Çalışmaları",
         "Bağımsız Geliştirme", "İstanbul  |  Devam ediyor")
bullet(doc, "Claude ve ChatGPT başta olmak üzere büyük dil modellerini teknik üretkenlik amacıyla sistematik biçimde kullanmaktayım.")
bullet(doc, "Yazılım geliştirme, hata kontrol listesi, proje analiz ve teknik dokümantasyon iyileştirme amaçlı yapılandırılmış promptlar hazırladım.")
bullet(doc, "Uzun, çok aşamalı master prompt yapıları geliştirerek AI destekli proje planlama süreçleri üzerine uygulamalı çalışmalar yaptım.")

# --- Çalışma 3 ---
heading2(doc, "Yazılım Mühendisliği Ders Projeleri",
         "İstanbul Arel Üniversitesi", "İstanbul  |  Devam ediyor")
bullet(doc, "Yazılım Mühendisliği dersi kapsamında gereksinim analizi, kullanım senaryosu modelleme, prototip tasarımı ve teknik dokümantasyon çalışmaları yürüttüm.")
bullet(doc, "Java ve SQL tabanlı temel yazılım geliştirme projeleri geliştirdim.")
bullet(doc, "Nesne yönelimli programlama, veri yapıları ve algoritma konularında uygulamalı çalışmalar gerçekleştirdim.")

# ═══════════════════════════════════════════════════════════════════════════
#  4. TEKNİK PROJELER
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "Teknik Projeler")

# ── Proje 1 ──
heading2(doc, "VARIANT-GNN — Sağlıkta Yapay Zekâ Projesi")
p_link = doc.add_paragraph()
rl = p_link.add_run("GitHub: ")
set_font(rl, size=9.5, bold=True, color=MID_GRAY)
rv = p_link.add_run("github.com/msgxr/VARIANT-GNN")
set_font(rv, size=9.5, color=ACCENT)
para_spacing(p_link, before=0, after=2)

body(doc,
     "Genetik varyant sınıflandırması üzerine geliştirilmekte olan, Graph Neural Network "
     "mimarisini temel alan bir yapay zekâ projesidir. Açıklanabilir yapay zekâ yöntemleriyle "
     "model çıktılarının yorumlanması ve klinik karar destek süreçleri için anlamlı çıktılar "
     "üretilmesi hedeflenmektedir.")
tag_line(doc, "Teknolojiler: Python · Graph Neural Networks · Streamlit · Jupyter Notebook · Git & GitHub")

bullet(doc, "GNN tabanlı model geliştirme süreci üzerine araştırma ve uygulama çalışmaları yürüttüm.")
bullet(doc, "Streamlit tabanlı demo ve görselleştirme arayüzü üzerine çalışmalar gerçekleştirdim.")
bullet(doc, "Teknik dokümantasyon hazırladım; hata analizi ve hata ayıklama süreçlerinde aktif rol aldım.")
bullet(doc, "GitHub repo yapısını ve dosya organizasyonunu düzenledim; sürüm yönetimi süreçlerini yönettim.")
bullet(doc, "Açıklanabilir yapay zekâ (XAI) çıktılarının yorumlanması üzerine analiz çalışmaları yaptım.")

p = doc.add_paragraph()
r = p.add_run("Geliştirilen Beceriler: ")
set_font(r, size=9.5, bold=True, color=MID_GRAY)
r2 = p.add_run("GNN mimarileri · Sağlık verisi analizi · Teknik proje yönetimi · Streamlit · XAI")
set_font(r2, size=9.5, color=BLACK)
para_spacing(p, before=1, after=6)

# ── Proje 2 ──
heading2(doc, "Prompt Engineering & AI Workflow Çalışmaları")
body(doc,
     "Büyük dil modellerini yazılım geliştirme, teknik analiz ve proje planlama süreçlerine "
     "entegre eden yapılandırılmış prompt sistemi ve AI iş akışları üzerine yürütülen bağımsız "
     "çalışma serisidir.")
tag_line(doc, "Teknolojiler: Claude API · ChatGPT · n8n · Python · VS Code · PowerShell")
bullet(doc, "Teknik şartnameye uygun yapılandırılmış master prompt sistemleri geliştirdim.")
bullet(doc, "Yazılım projelerinde hata tespiti ve hata kontrol listesi oluşturma amacıyla özel promptlar hazırladım.")
bullet(doc, "Kod inceleme, dokümantasyon iyileştirme ve proje analiz amacıyla AI iş akışları tasarladım.")
bullet(doc, "n8n ile otomasyon iş akışı prototipleri üzerine fikir ve tasarım çalışmaları yürüttüm.")
p = doc.add_paragraph()
r = p.add_run("Geliştirilen Beceriler: ")
set_font(r, size=9.5, bold=True, color=MID_GRAY)
r2 = p.add_run("LLM prompt mühendisliği · AI destekli yazılım geliştirme · Teknik iş akışı tasarımı")
set_font(r2, size=9.5, color=BLACK)
para_spacing(p, before=1, after=6)

# ── Proje 3 ──
heading2(doc, "Java / SQL Tabanlı Yazılım Çalışmaları")
body(doc,
     "Üniversite ders süreçleri kapsamında Java ve SQL kullanılarak geliştirilen temel yazılım "
     "projeleri ve uygulamalı programlama çalışmaları.")
tag_line(doc, "Teknolojiler: Java · SQL · OOP Prensipleri · Temel Veritabanı Tasarımı")
bullet(doc, "Nesne yönelimli tasarım prensipleriyle Java uygulamaları geliştirdim.")
bullet(doc, "SQL tabanlı veritabanı sorgulama, ilişkisel şema tasarımı ve veri yönetimi çalışmaları yürüttüm.")
bullet(doc, "Yazılım Mühendisliği dersi kapsamında gereksinim analizi, modelleme ve dokümantasyon çalışmaları gerçekleştirdim.")
p = doc.add_paragraph()
r = p.add_run("Geliştirilen Beceriler: ")
set_font(r, size=9.5, bold=True, color=MID_GRAY)
r2 = p.add_run("OOP prensipleri · Veritabanı yönetimi · Yazılım tasarımı · Teknik dokümantasyon")
set_font(r2, size=9.5, color=BLACK)
para_spacing(p, before=1, after=6)

# ── Proje 4 ──
heading2(doc, "Yapay Zekâ Destekli Finansal Analiz & Otomasyon Fikir Çalışması")
body(doc,
     "BIST100 ve finansal veri akışları üzerine yapay zekâ destekli analiz yaklaşımları, "
     "otomasyon iş akışları ve veri işleme senaryoları üzerine yürütülen araştırma ve "
     "fikir geliştirme çalışması.")
tag_line(doc, "Teknolojiler: Python · n8n · Huawei Cloud · LLM Araçları")
bullet(doc, "Finansal veri akışlarında yapay zekâ destekli analiz senaryoları araştırdım.")
bullet(doc, "n8n ve Huawei Cloud servisleriyle otomasyon entegrasyonu prototipleri üzerine tasarım çalışmaları yaptım.")
bullet(doc, "Veri odaklı karar destek sistemleri için kavramsal mimari geliştirdim.")
p = doc.add_paragraph()
r = p.add_run("Geliştirilen Beceriler: ")
set_font(r, size=9.5, bold=True, color=MID_GRAY)
r2 = p.add_run("Finansal veri analizi · Otomasyon tasarımı · Bulut servis entegrasyonu · Sistem mimarisi")
set_font(r2, size=9.5, color=BLACK)
para_spacing(p, before=1, after=6)

# ═══════════════════════════════════════════════════════════════════════════
#  5. EĞİTİM
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "Eğitim")

heading2(doc, "İstanbul Arel Üniversitesi",
         "Bilgisayar Mühendisliği — Lisans", "İstanbul, Türkiye  |  Devam ediyor")

p = doc.add_paragraph()
r = p.add_run("İlgili Dersler: ")
set_font(r, size=10, bold=True, color=MID_GRAY)
r2 = p.add_run(
    "Veri Yapıları ve Algoritmalar · Nesne Yönelimli Programlama · "
    "Veritabanı Yönetim Sistemleri · Yazılım Mühendisliği · "
    "Programlama Temelleri · Matematik · Fizik"
)
set_font(r2, size=10, color=BLACK)
para_spacing(p, before=2, after=4)

# ═══════════════════════════════════════════════════════════════════════════
#  6. BAŞARILAR VE ETKİNLİKLER
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "Başarılar ve Etkinlikler")

basarilar = [
    ("Sağlıkta Yapay Zekâ Projesi",
     "VARIANT-GNN kapsamında GNN tabanlı genetik varyant sınıflandırması üzerine aktif proje geliştirme çalışmaları yürütmekteyim."),
    ("GitHub Teknik Proje Geliştirme",
     "github.com/msgxr üzerinde proje repo yönetimi, teknik dokümantasyon ve sürüm yönetimi çalışmaları gerçekleştirmekteyim."),
    ("Prompt Engineering Uygulamaları",
     "Büyük dil modellerini teknik üretkenlik amacıyla sistematik biçimde kullanan yapılandırılmış prompt sistemleri geliştirdim."),
    ("Yazılım Mühendisliği Projeleri",
     "Ders kapsamındaki projelerde gereksinim analizi, yazılım modelleme, prototip geliştirme ve teknik dokümantasyon çalışmaları yürüttüm."),
    ("Teknoloji Yarışmaları & Hackathonlar",
     "Yapay zekâ, yazılım geliştirme ve teknoloji alanlarındaki yarışma ve etkinliklere aktif ilgi göstermekte, başvuru ve katılım süreçlerini takip etmekteyim."),
    ("Otomasyon & Bulut Servisleri",
     "n8n ve Huawei Cloud ile otomasyon ve bulut altyapısı üzerine araştırma ve uygulama çalışmaları yürüttüm."),
]

for baslik, aciklama in basarilar:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{baslik}: ")
    set_font(r1, size=10, bold=True, color=DARK_NAVY)
    r2 = p.add_run(aciklama)
    set_font(r2, size=10, color=BLACK)
    para_spacing(p, before=2, after=1)
    p.paragraph_format.left_indent = Inches(0.2)

# ═══════════════════════════════════════════════════════════════════════════
#  7. DİLLER
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "Diller")

diller = [
    ("Türkçe", "Anadil"),
    ("İngilizce", "B1 — Teknik okuma ve dokümantasyon kullanımı"),
]
for dil, seviye in diller:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{dil}: ")
    set_font(r1, size=10, bold=True, color=DARK_NAVY)
    r2 = p.add_run(seviye)
    set_font(r2, size=10, color=BLACK)
    para_spacing(p, before=2, after=1)
    p.paragraph_format.left_indent = Inches(0.2)

# ═══════════════════════════════════════════════════════════════════════════
#  8. LİNKEDİN BAŞLIĞI & KISA ÖZET
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "LinkedIn Başlığı")

p = doc.add_paragraph()
r = p.add_run(
    "Bilgisayar Mühendisliği Öğrencisi  |  Yapay Zekâ & Makine Öğrenmesi  |  "
    "Prompt Engineering  |  VARIANT-GNN  |  GitHub: msgxr"
)
set_font(r, size=10, italic=True, color=MID_GRAY)
para_spacing(p, before=2, after=6)

heading1(doc, "Kısa Profesyonel Özet")

kisa_ozet = (
    "İstanbul Arel Üniversitesi Bilgisayar Mühendisliği öğrencisiyim. "
    "Yapay zekâ, makine öğrenmesi ve prompt engineering alanlarında uygulamalı proje çalışmaları yürütmekteyim. "
    "Sağlık alanında GNN tabanlı genetik varyant sınıflandırması projesi olan VARIANT-GNN üzerinde aktif olarak "
    "geliştirme, teknik dokümantasyon ve hata analizi çalışmaları gerçekleştirmekteyim. "
    "Büyük dil modellerini yazılım geliştirme ve proje planlama süreçlerine entegre eden yapılandırılmış "
    "prompt sistemleri geliştiriyorum. "
    "GitHub üzerinde aktif proje yönetimiyle sistematik ve belgelenmiş bir çalışma anlayışı benimsemekteyim."
)
body(doc, kisa_ozet)

# ═══════════════════════════════════════════════════════════════════════════
#  KAYDET
# ═══════════════════════════════════════════════════════════════════════════

output_path = "muhammedsinagun_cv_tr.docx"
doc.save(output_path)
print(f"CV oluşturuldu: {output_path}")
